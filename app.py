import io
import gc
import cv2
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import numpy as np
from PIL import Image
import streamlit as st

# ---------------------------------------------------------
# CẤU HÌNH TRANG & GIAO DIỆN
# ---------------------------------------------------------
st.set_page_config(
    page_title="Hệ Thống Xử Lý & Dàn Trang Giấy Tờ",
    layout="wide",
    page_icon="🖨️",
)

st.markdown(
    """
    <style>
    .main-header {
        font-size: 1.8rem;
        font-weight: 700;
        color: #1E3A8A;
        margin-bottom: 0.2rem;
    }
    .stDownloadButton button {
        width: 100%;
        background-color: #16A34A !important;
        color: white !important;
        font-weight: bold !important;
        font-size: 1.05rem !important;
        border-radius: 8px !important;
        padding: 0.75rem 1rem !important;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    .stDownloadButton button:hover {
        background-color: #15803D !important;
    }
    </style>
""",
    unsafe_allow_html=True,
)

st.markdown('<div class="main-header">🖨️ Xử Lý & Dàn Trang Giấy Tờ Super Fast</div>', unsafe_allow_html=True)

# Quản lý Session State
if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = 0
if "swap_dict" not in st.session_state:
    st.session_state["swap_dict"] = {}

def clear_all_files():
    st.session_state["uploader_key"] += 1
    st.session_state["swap_dict"] = {}

def toggle_swap_pair(pair_index):
    current_state = st.session_state["swap_dict"].get(pair_index, False)
    st.session_state["swap_dict"][pair_index] = not current_state

# ---------------------------------------------------------
# THUẬT TOÁN XỬ LÝ ẢNH CHUNG
# ---------------------------------------------------------
def optimize_image_size(pil_img, max_dim=1600):
    w, h = pil_img.size
    if max(w, h) > max_dim:
        scale = max_dim / float(max(w, h))
        return pil_img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
    return pil_img

def crop_card_fast(image_np):
    h_img, w_img, _ = image_np.shape
    gray = cv2.cvtColor(image_np, cv2.COLOR_RGB2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thresh = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    contours, _ = cv2.findContours(thresh, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)

    best_rect, max_area = None, 0
    for cnt in contours:
        area = cv2.contourArea(cnt)
        if area > (w_img * h_img * 0.10):
            x, y, w, h = cv2.boundingRect(cnt)
            aspect_ratio = float(w) / h
            if 1.2 <= aspect_ratio <= 1.95 and (w < w_img * 0.98):
                if area > max_area:
                    max_area = area
                    best_rect = (x, y, w, h)

    if best_rect:
        x, y, w, h = best_rect
        return image_np[max(0, y - 4):min(h_img, y + h + 8), max(0, x - 4):min(w_img, x + w + 8)]
    return image_np

def add_card_row_to_word(doc, pil_front, pil_back):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    buf_front, buf_back = io.BytesIO(), io.BytesIO()
    pil_front.save(buf_front, format="PNG", optimize=True)
    pil_back.save(buf_back, format="PNG", optimize=True)
    buf_front.seek(0); buf_back.seek(0)

    cell_f = table.cell(0, 0)
    p_f = cell_f.paragraphs[0]
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.add_run().add_picture(buf_front, width=Inches(3.35))

    cell_b = table.cell(0, 1)
    p_b = cell_b.paragraphs[0]
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_b.add_run().add_picture(buf_back, width=Inches(3.35))

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(20)

def create_multi_docx(card_pairs):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    for i, (pil_front, pil_back) in enumerate(card_pairs):
        if i > 0 and i % 2 == 0:
            doc.add_page_break()
        add_card_row_to_word(doc, pil_front, pil_back)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_single_cropped_docx(pil_cropped_img):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    buf = io.BytesIO()
    pil_cropped_img.save(buf, format="PNG", optimize=True)
    buf.seek(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(3.5))

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_a4_canvas_horizontal(card_pairs_chunk):
    a4_w, a4_h = 1240, 1754
    canvas = Image.new("RGB", (a4_w, a4_h), "white")
    card_w, spacing_x = 540, 40
    start_x = (a4_w - (card_w * 2 + spacing_x)) // 2
    y_positions = [180, 920]

    for idx, (pil_front, pil_back) in enumerate(card_pairs_chunk):
        if idx >= 2: break
        y_pos = y_positions[idx]
        ratio_f = card_w / pil_front.width
        f_resized = pil_front.resize((card_w, int(pil_front.height * ratio_f)), Image.Resampling.LANCZOS)
        ratio_b = card_w / pil_back.width
        b_resized = pil_back.resize((card_w, int(pil_back.height * ratio_b)), Image.Resampling.LANCZOS)

        canvas.paste(f_resized, (start_x, y_pos))
        canvas.paste(b_resized, (start_x + card_w + spacing_x, y_pos))

    return canvas

# ---------------------------------------------------------
# THUẬT TOÁN CẮT KHUNG VNeID
# ---------------------------------------------------------
def crop_vneid_combined_block(pil_img):
    img_np = np.array(pil_img)
    h_img, w_img, _ = img_np.shape
    gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thresh = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    contours, _ = cv2.findContours(thresh, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    
    card_boxes = []
    for cnt in contours:
        area = cv2.contourArea(cnt)
        if area > (w_img * h_img * 0.05):
            x, y, w, h = cv2.boundingRect(cnt)
            aspect_ratio = float(w) / h
            if 1.2 <= aspect_ratio <= 1.95 and w < w_img * 0.98:
                card_boxes.append((x, y, w, h))

    filtered_boxes = []
    card_boxes = sorted(card_boxes, key=lambda b: b[1])
    for box in card_boxes:
        if not any(abs(box[1] - f[1]) < 30 for f in filtered_boxes):
            filtered_boxes.append(box)

    if len(filtered_boxes) >= 2:
        b1, b2 = filtered_boxes[0], filtered_boxes[1]
        min_x = min(b1[0], b2[0])
        max_x = max(b1[0] + b1[2], b2[0] + b2[2])
        min_y = min(b1[1], b2[1])
        max_y = max(b1[1] + b1[3], b2[1] + b2[3])

        crop_np = img_np[
            max(0, min_y - 6) : min(h_img, max_y + 14), 
            max(0, min_x - 6) : min(w_img, max_x + 6)
        ]
        return Image.fromarray(crop_np)

    y1, y2 = int(h_img * 0.11), int(h_img * 0.67)
    x1, x2 = int(w_img * 0.035), int(w_img * 0.965)
    
    crop_np = img_np[y1:y2, x1:x2]
    return Image.fromarray(crop_np)


# ---------------------------------------------------------
# KHỞI TẠO TABS GIAO DIỆN (ĐÃ ĐẶT TRƯỚC WITH TAB)
# ---------------------------------------------------------
tab1, tab2 = st.tabs(["🖨️ Dàn Trang A4 (Hàng Loạt)", "✂️ Cắt Ảnh Khung GPLX VNeID"])

# =========================================================
# TAB 1: DÀN TRANG A4 TỰ ĐỘNG
# =========================================================
with tab1:
    col_left, col_right = st.columns([1, 1], gap="medium")

    with col_left:
        st.subheader("1. Tải Ảnh & Xuất File")
        
        doc_type = st.radio(
            "📋 Loại giấy tờ:",
            ["CCCD", "GPLX"],
            horizontal=True
        )

        col_up_title, col_btn_clear = st.columns([0.65, 0.35])
        with col_up_title:
            st.write("📥 Tải lên tất cả ảnh:")
        with col_btn_clear:
            st.button("🧹 Làm mới", on_click=clear_all_files, use_container_width=True)

        uploaded_files = st.file_uploader(
            "Tải lên tất cả ảnh:",
            type=["jpg", "jpeg", "png"],
            accept_multiple_files=True,
            key=f"file_uploader_{st.session_state['uploader_key']}",
            label_visibility="collapsed"
        )

        download_area = st.container()

    with col_right:
        st.subheader("2. Xem Trước Bản In A4 & Tuỳ Chỉnh")
        preview_area = st.container()

    if uploaded_files:
        num_files = len(uploaded_files)

        # Tính toán số cặp (Lẻ 1 ảnh vẫn xử lý được)
        num_pairs = (num_files + 1) // 2
        card_pairs = []
        a4_canvases = []

        with col_left:
            progress_bar = st.progress(0)
            status_text = st.empty()

            for idx in range(num_pairs):
                status_text.text(f"⏳ Đang xử lý Bộ {idx + 1}/{num_pairs}...")
                
                i1 = idx * 2
                i2 = i1 + 1

                # Ảnh thứ 1
                raw_img1 = Image.open(uploaded_files[i1]).convert("RGB")
                opt_1 = optimize_image_size(raw_img1)
                crop_1 = crop_card_fast(np.array(opt_1))
                pil_f = Image.fromarray(crop_1)

                # Ảnh thứ 2 (Nếu thiếu ảnh thì bù nền trắng)
                if i2 < num_files:
                    raw_img2 = Image.open(uploaded_files[i2]).convert("RGB")
                    opt_2 = optimize_image_size(raw_img2)
                    crop_2 = crop_card_fast(np.array(opt_2))
                    pil_b = Image.fromarray(crop_2)
                    del raw_img2, opt_2, crop_2
                else:
                    pil_b = Image.new("RGB", pil_f.size, (255, 255, 255))

                if st.session_state["swap_dict"].get(idx, False):
                    pil_f, pil_b = pil_b, pil_f

                card_pairs.append((pil_f, pil_b))

                del raw_img1, opt_1, crop_1
                gc.collect()

                progress_bar.progress((idx + 1) / num_pairs)

            status_text.empty()
            progress_bar.empty()

        for i in range(0, len(card_pairs), 2):
            chunk = card_pairs[i:i + 2]
            canvas = create_a4_canvas_horizontal(chunk)
            a4_canvases.append(canvas)

        docx_io = create_multi_docx(card_pairs)

        with download_area:
            if num_files % 2 != 0:
                st.warning("⚠️ Số lượng ảnh lẻ, hệ thống đã tự ghép ảnh trống làm mặt còn lại.")
            st.success(f"⚡ Đã ghép xong **{len(card_pairs)} bộ**!")

            st.download_button(
                label=f"📝 TẢI FILE WORD NGAY (.docx)",
                data=docx_io,
                file_name=f"{doc_type}_In_A4.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            
            img_io = io.BytesIO()
            a4_canvases[0].save(img_io, format="PNG", optimize=True)
            img_io.seek(0)

            st.download_button(
                label="🖼️ TẢI ẢNH PNG TRANG 1",
                data=img_io,
                file_name=f"{doc_type}_Trang1.png",
                mime="image/png",
            )

        with preview_area:
            st.write("🔄 **Tuỳ chỉnh vị trí cho riêng từng bộ (nếu bị ngược):**")
            btn_cols = st.columns(min(len(card_pairs), 4))
            for p_idx in range(len(card_pairs)):
                is_swapped = st.session_state["swap_dict"].get(p_idx, False)
                btn_label = f"🔄 Đảo Bộ {p_idx + 1}" + (" (Đã đảo)" if is_swapped else "")
                btn_cols[p_idx % 4].button(
                    btn_label, 
                    key=f"btn_swap_{p_idx}", 
                    on_click=toggle_swap_pair, 
                    args=(p_idx,),
                    use_container_width=True
                )

            st.markdown("---")

            if len(a4_canvases) > 1:
                tabs = st.tabs([f"Trang #{i + 1}" for i in range(len(a4_canvases))])
                for i, tab in enumerate(tabs):
                    with tab:
                        st.image(a4_canvases[i], use_container_width=True)
            else:
                st.image(a4_canvases[0], use_container_width=True)


# =========================================================
# TAB 2: CẮT KHUNG VNeID CHUẨN MẪU
# =========================================================
with tab2:
    st.subheader("✂️ Cắt Trọn Khung Bằng Lái Xe GPLX VNeID")
    st.caption("Tải lên ảnh màn hình VNeID, hệ thống sẽ tự động cắt khung chứa 2 thẻ GPLX.")

    col_vneid_up, col_vneid_res = st.columns([1, 1], gap="medium")

    with col_vneid_up:
        uploaded_vneid = st.file_uploader(
            "📥 Tải ảnh chụp màn hình GPLX VNeID:",
            type=["jpg", "jpeg", "png"],
            key="vneid_uploader"
        )
        if uploaded_vneid:
            raw_vneid_img = Image.open(uploaded_vneid).convert("RGB")
            st.image(raw_vneid_img, caption="Ảnh gốc VNeID", use_container_width=True)

    with col_vneid_res:
        if uploaded_vneid:
            with st.spinner("⏳ Đang nhận diện và cắt khung ảnh..."):
                cropped_gplx_img = crop_vneid_combined_block(raw_vneid_img)

                buf = io.BytesIO()
                cropped_gplx_img.save(buf, format="PNG", optimize=True)
                buf_png = buf.getvalue()

                docx_cropped_io = create_single_cropped_docx(cropped_gplx_img)

                st.markdown("### 📥 Tải về kết quả:")
                
                col_dl_word, col_dl_png = st.columns(2)
                with col_dl_word:
                    st.download_button(
                        label="📝 TẢI FILE WORD (.docx)",
                        data=docx_cropped_io,
                        file_name="GPLX_VNeID_Cropped.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with col_dl_png:
                    st.download_button(
                        label="🖼️ TẢI ẢNH PNG (.png)",
                        data=buf_png,
                        file_name="GPLX_VNeID_Cropped.png",
                        mime="image/png",
                        use_container_width=True
                    )

                st.markdown("---")
                st.markdown("### 👁️ Xem trước ảnh đã cắt:")
                st.image(cropped_gplx_img, caption="Khung ảnh GPLX 2 mặt", use_container_width=True)
